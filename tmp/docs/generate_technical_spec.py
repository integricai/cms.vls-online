"""Generate VLS Online platform technical specification Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[2] / "output" / "doc" / "VLS-Online-Platform-Technical-Specification.docx"


def set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x0D, 0x1F, 0x3C)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_paragraph("(Update table of contents in Word: right-click and choose Update Field.)")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    set_margins(doc)
    style_doc(doc)

    # --- Cover ---
    for _ in range(6):
        doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("VLS Online Platform\nTechnical Specification")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x0D, 0x1F, 0x3C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Storyblok CMS + Next.js + Cloudflare + Zenler LMS\nMarketing Site Replatform")
    sr.font.size = Pt(14)

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Version 1.0\n{date.today().strftime('%d %B %Y')}\nvls-online.com")
    mr.font.size = Pt(12)

    doc.add_page_break()

    # --- Document control ---
    doc.add_heading("Document Control", level=1)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Document title", "VLS Online Platform Technical Specification"],
            ["Project", "vls-online.com marketing replatform"],
            ["Version", "1.0"],
            ["Status", "Draft for review"],
            ["Date", date.today().strftime("%Y-%m-%d")],
            ["Audience", "Engineering, product, content operations"],
            ["Supersedes", "cms-v2 inject-based architecture (marketing layer)"],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("Revision History", level=2)
    add_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [["1.0", date.today().strftime("%Y-%m-%d"), "VLS Engineering", "Initial technical specification"]],
        [0.8, 1.0, 1.2, 2.5],
    )

    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # --- 1 Executive Summary ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document defines the technical specification for replatforming the VLS Online public "
        "marketing website from the current cms-v2 HTML-inject model hosted on Zenler to a modern "
        "architecture comprising Storyblok as the content management system, Next.js as the rendering "
        "platform, Cloudflare as the edge and routing layer, and Zenler retained exclusively as the "
        "learning management system (LMS) for enrolled students."
    )
    doc.add_paragraph(
        "The primary business drivers are: enabling non-technical editors to create and publish landing "
        "pages without developer involvement; improving SEO and Core Web Vitals; supporting multi-market "
        "translated content; implementing geo-aware pricing via Stripe; and eliminating fragile "
        "copy-paste HTML workflows that cause production UI defects."
    )

    doc.add_heading("1.1 Goals", level=2)
    add_bullets(doc, [
        "Replace cms-v2 marketing content authoring with Storyblok Visual Editor.",
        "Serve all public marketing pages from a first-party Next.js application.",
        "Retain Zenler for video delivery, curriculum, student login, progress, and certificates.",
        "Implement Stripe geo-pricing with automatic Zenler enrollment after successful payment.",
        "Support internationalization (i18n) for priority markets.",
        "Achieve full control of structured data (schema.org), sitemaps, and page metadata.",
        "Decommission HTML inject scripts and publish-* snippet endpoints after migration.",
    ])

    doc.add_heading("1.2 Non-Goals", level=2)
    add_bullets(doc, [
        "Building a custom LMS or custom video streaming platform in Phase 1.",
        "Replacing Zenler video hosting with AWS S3 or Google Cloud Storage.",
        "Migrating historical student progress or accounts out of Zenler.",
        "Retaining cms-v2 React admin editors post cutover.",
    ])

    doc.add_heading("1.3 Success Criteria", level=2)
    add_table(doc,
        ["Metric", "Target"],
        [
            ["Lighthouse Performance (mobile)", ">= 80 on course landing pages"],
            ["Lighthouse SEO", ">= 95"],
            ["Editor publish time", "< 2 minutes without developer (after training)"],
            ["Time to first contentful paint", "< 2.0s on 4G (cached ISR pages)"],
            ["Geo price resolution", "< 100ms API p95"],
            ["Enrollment after payment", "< 60s from Stripe webhook to Zenler access"],
        ],
        [3.0, 3.5],
    )

    # --- 2 Current State ---
    doc.add_heading("2. Current State Assessment", level=1)
    doc.add_paragraph(
        "The existing cms-v2 system (repository: cms.vls-online) is a custom CMS comprising a React "
        "admin application and Express API deployed on Vercel with Neon PostgreSQL. Marketing content "
        "is authored as JSON in Neon, rendered to self-contained HTML/JS snippets via generateHtml.ts "
        "modules, and pasted into Zenler page blocks on vls-online.com."
    )

    doc.add_heading("2.1 Current Architecture Components", level=2)
    add_table(doc,
        ["Component", "Technology", "Role"],
        [
            ["cms-v2 admin", "React 18, Vite, Tailwind", "40+ component editors for marketing snippets"],
            ["cms-v2 API", "Node.js, Express, TypeScript", "Content CRUD, blog SSR, Stripe, Zenler sync"],
            ["Database", "Neon PostgreSQL", "cms_content, courses, blog, payments, books"],
            ["Public site host", "Zenler (vls.newzenler.com)", "Page hosting, LMS, video, checkout (partial)"],
            ["Payments", "Stripe Checkout (partial)", "Payment cards bypass Zenler billing for some courses"],
            ["Edge", "Cloudflare Turnstile only", "CAPTCHA on forms; limited routing"],
        ],
        [1.5, 1.8, 3.2],
    )

    doc.add_heading("2.2 Known Limitations", level=2)
    add_table(doc,
        ["Area", "Issue", "Impact"],
        [
            ["Content workflow", "Developers required for layout/component changes", "High operational cost, slow campaigns"],
            ["Publishing", "Copy-paste HTML injects into Zenler", "Frequent UI bugs from editor mistakes"],
            ["Page flexibility", "Fixed cms-v2 component model", "No ad-hoc landing pages"],
            ["SEO", "Zenler strips injected schema; heavy third-party JS", "Poor search visibility"],
            ["Performance", "No lazy loading; slow first content load", "Poor conversion and Core Web Vitals"],
            ["Navigation", "Zenler menu lacks reliable 2nd-level items", "Poor UX; cms-v2 header hacks required"],
            ["Blog", "Requires inject scripts to hide Zenler chrome", "Fragile, hard to maintain"],
            ["i18n", "Not supported in cms-v2", "Cannot serve translated markets"],
            ["Pricing", "Single currency per course; Zenler has no geo pricing", "Revenue leakage in international markets"],
        ],
        [1.2, 2.5, 2.8],
    )

    doc.add_heading("2.3 Retained Assets from cms-v2", level=2)
    add_bullets(doc, [
        "Stripe checkout session creation (cms-v2/src/services/stripeCheckout.ts).",
        "Stripe webhook verification and payment order lifecycle (cms-v2/src/routes/payments.ts).",
        "Zenler course catalog sync (cms-v2/src/services/zenlerCourseService.ts).",
        "Neon database schema for courses, payment cards, orders, books.",
        "Form submission and Turnstile integration patterns.",
        "Blog content and assets (subject to migration into Storyblok).",
    ])

    # --- 3 Target Architecture ---
    doc.add_heading("3. Target Architecture", level=1)
    doc.add_paragraph(
        "The target system separates concerns into four layers: Storyblok (content), Next.js (presentation), "
        "Cloudflare (edge routing, geo, cache, security), and a slim backend API evolved from cms-v2 "
        "(commerce and integrations). Zenler operates on a dedicated subdomain for enrolled students only."
    )

    doc.add_heading("3.1 System Context Diagram (Logical)", level=2)
    add_code(doc, """
[Visitor] --> [Cloudflare DNS/WAF/Worker]
                  |
      +-----------+-----------+
      |                       |
 [Next.js Site]          [Zenler LMS]
 (vls-online.com)    (learn.vls-online.com)
      |                       |
      +--> [Storyblok CMS]    +--> [Video/Curriculum]
      |
      +--> [Backend API] --> [Neon PostgreSQL]
                  |
                  +--> [Stripe]
                  +--> [Zenler Enrollment API]
""")

    doc.add_heading("3.2 Component Responsibilities", level=2)
    add_table(doc,
        ["Layer", "Product", "Responsibility"],
        [
            ["CMS", "Storyblok", "Page content, blog, translations, media, blok schema, editorial workflow"],
            ["Frontend", "Next.js 15 (App Router)", "SSR/SSG/ISR, SEO, schema.org, i18n routing, checkout UI"],
            ["Frontend hosting", "Vercel", "Build, deploy, preview environments, ISR"],
            ["Edge", "Cloudflare", "DNS, SSL, CDN cache, WAF, Workers routing, CF-IPCountry, Turnstile"],
            ["Backend", "Express API (cms-v2 evolved)", "Stripe, webhooks, geo prices, Zenler sync/enrollment, forms"],
            ["Database", "Neon PostgreSQL", "Courses, regional prices, orders, books, activity logs"],
            ["LMS", "Zenler", "Login, course player, curriculum, progress, certificates, video hosting"],
            ["Payments", "Stripe", "Multi-currency checkout, webhooks, customer receipts"],
        ],
        [1.0, 1.5, 4.0],
    )

    doc.add_heading("3.3 Design Principles", level=2)
    add_numbered(doc, [
        "Content and presentation decoupling: Storyblok stores content; Next.js owns rendering and performance.",
        "Commerce source of truth: Stripe Price IDs and regional price tables drive displayed and charged amounts.",
        "Zenler scope minimization: Zenler is used only post-purchase for learning delivery.",
        "No HTML inject publishing: all marketing output is rendered by Next.js from structured content.",
        "Progressive migration: route migrated paths to Next.js; maintain 301 redirects for SEO continuity.",
        "Editor self-service: new landing pages creatable from Storyblok without code changes.",
    ])

    # --- 4 Domain and Routing ---
    doc.add_heading("4. Domain and Routing Specification", level=1)

    doc.add_heading("4.1 Domain Map", level=2)
    add_table(doc,
        ["Hostname", "Origin", "Purpose"],
        [
            ["vls-online.com", "Vercel (Next.js)", "Public marketing website"],
            ["www.vls-online.com", "301 redirect", "Redirect to apex vls-online.com"],
            ["learn.vls-online.com", "Zenler custom domain", "Student LMS portal"],
            ["api.vls-online.com", "Vercel (Express API)", "Backend services"],
            ["preview.vls-online.com", "Vercel preview", "Storyblok draft preview (optional)"],
        ],
        [2.0, 2.0, 2.5],
    )

    doc.add_heading("4.2 Marketing URL Routes (Next.js)", level=2)
    add_table(doc,
        ["Path Pattern", "Content Source", "Render Mode"],
        [
            ["/", "Storyblok: pages/home", "ISR (revalidate: 60s)"],
            ["/[locale]", "Storyblok: pages/home (localized)", "ISR"],
            ["/[locale]/courses/[slug]", "Storyblok: courses/{slug}", "ISR + dynamic pricing"],
            ["/[locale]/blog", "Storyblok: blog listing", "ISR"],
            ["/[locale]/blog/[slug]", "Storyblok: blog/posts/{slug}", "ISR"],
            ["/[locale]/[...slug]", "Storyblok: flexible pages", "ISR"],
            ["/checkout/[courseSlug]", "Next.js + API + Stripe", "Dynamic (SSR)"],
            ["/payment-success", "Next.js", "Dynamic"],
            ["/payment-cancelled", "Next.js", "Dynamic"],
            ["/sitemap.xml", "Generated", "Dynamic"],
            ["/robots.txt", "Static config", "Static"],
        ],
        [2.2, 2.3, 2.0],
    )

    doc.add_heading("4.3 LMS URL Routes (Zenler on learn.vls-online.com)", level=2)
    add_table(doc,
        ["Path Pattern", "Purpose"],
        [
            ["/login", "Student authentication"],
            ["/signup", "Student registration (if enabled)"],
            ["/dashboard", "Student home"],
            ["/my-courses", "Enrolled courses list"],
            ["/course_player/*", "Video and lesson player"],
            ["/lessons/*", "Individual lesson access"],
            ["/programs/*", "Program bundles (if used)"],
        ],
        [3.0, 3.5],
    )

    doc.add_heading("4.4 Cloudflare Worker Routing (Optional Path-Based Fallback)", level=2)
    doc.add_paragraph(
        "Primary recommendation is subdomain separation (Section 4.1). If same-domain LMS paths are "
        "required during transition, extend workers/worker.js with explicit routing rules."
    )
    add_code(doc, """
const MARKETING_HOST = 'vls-online.com';
const LMS_HOST = 'learn.vls-online.com';

export default {
  async fetch(request, env) {
    const url = new URL(request.url);
    const country = request.cf?.country ?? 'GB';

    // Apex domain: all traffic to Next.js (recommended)
    if (url.hostname === MARKETING_HOST) {
      const req = new Request(request);
      req.headers.set('x-vls-country', country);
      return fetch(env.MARKETING_ORIGIN + url.pathname + url.search, req);
    }

    // LMS subdomain: proxy to Zenler
    if (url.hostname === LMS_HOST) {
      return fetch(env.ZENLER_ORIGIN + url.pathname + url.search, request);
    }

    return new Response('Not found', { status: 404 });
  }
};
""")

    doc.add_heading("4.5 SEO Redirect Map (Migration)", level=2)
    doc.add_paragraph(
        "All legacy Zenler marketing URLs on vls-online.com must 301 redirect to equivalent Next.js routes. "
        "Maintain a redirect registry in Cloudflare Bulk Redirects or Next.js middleware."
    )
    add_table(doc,
        ["Legacy URL", "New URL", "Notes"],
        [
            ["/courses/{slug} (Zenler sales page)", "/en/courses/{slug}", "Preserve slug where possible"],
            ["/blog/{topic}/{slug}", "/en/blog/{slug}", "Drop Zenler topic segment"],
            ["/blog/{slug}", "/en/blog/{slug}", "Direct mapping"],
            ["/login (on apex)", "https://learn.vls-online.com/login", "Cross-subdomain redirect"],
            ["/course_player/* (on apex)", "https://learn.vls-online.com/course_player/*", "Cross-subdomain"],
        ],
        [2.5, 2.5, 1.5],
    )

    # --- 5 Storyblok ---
    doc.add_heading("5. Storyblok CMS Specification", level=1)

    doc.add_heading("5.1 Storyblok Plan Requirements", level=2)
    add_table(doc,
        ["Feature", "Requirement"],
        [
            ["Visual Editor", "Required for all page types"],
            ["Nested bloks", "Required for flexible landing pages"],
            ["Translatable fields", "Required for i18n markets"],
            ["Webhooks", "Required for ISR revalidation on publish"],
            ["Roles and workflows", "Editor, Admin, Developer roles"],
            ["Asset manager", "Images, PDFs, icons with folders"],
            ["Preview URLs", "Vercel preview deployment integration"],
        ],
        [2.5, 3.5],
    )

    doc.add_heading("5.2 Space Folder Structure", level=2)
    add_code(doc, """
global/
  site-config
  header
  footer
  announcement-bar

pages/
  home
  about
  contact
  legal/privacy-policy
  legal/terms-and-conditions

courses/
  fa1
  fa2
  ...

blog/
  posts/
    {slug}
  topics/
    {slug}          (optional category landing pages)

landing/
  {campaign-slug}   (editors create freely)

locales/
  en/...            (if using folder-per-locale strategy)
  ar/...
""")

    doc.add_heading("5.3 Content Type: page (Flexible Landing Page)", level=2)
    add_table(doc,
        ["Field", "Type", "Required", "Description"],
        [
            ["title", "Text", "Yes", "Internal and display title"],
            ["slug", "Text (slug)", "Yes", "URL segment"],
            ["seo_title", "Text", "No", "Meta title override"],
            ["seo_description", "Text", "No", "Meta description"],
            ["og_image", "Asset", "No", "Open Graph image"],
            ["no_index", "Boolean", "No", "Exclude from search engines"],
            ["body", "Blocks", "Yes", "Nested section bloks (see 5.5)"],
        ],
        [1.3, 1.0, 0.7, 3.5],
    )

    doc.add_heading("5.4 Content Type: course_page (Structured Course Landing)", level=2)
    add_table(doc,
        ["Field", "Type", "Required", "Description"],
        [
            ["title", "Text", "Yes", "Course display name"],
            ["slug", "Text", "Yes", "URL: /courses/{slug}"],
            ["zenler_course_id", "Text", "Yes", "Links to Zenler catalog and enrollment"],
            ["course_code", "Text", "No", "e.g. FA1, ACCA"],
            ["hero", "Block: section_hero", "Yes", "Above-the-fold content"],
            ["outcomes", "Blocks: section_richtext or list", "No", "Learning outcomes"],
            ["curriculum_summary", "Blocks", "No", "Static curriculum overview (not live LMS data)"],
            ["pricing", "Block: course_price_block", "Yes", "Dynamic price display config"],
            ["testimonials", "Block: section_testimonials", "No", "Social proof"],
            ["faq", "Block: section_faq", "No", "Course FAQ"],
            ["cta", "Block: section_cta", "No", "Final enrollment CTA"],
            ["schema", "Block: schema_course", "No", "JSON-LD overrides"],
            ["seo_*", "Various", "No", "Standard SEO fields"],
        ],
        [1.4, 1.1, 0.6, 3.4],
    )

    doc.add_heading("5.5 Content Type: blog_post", level=2)
    add_table(doc,
        ["Field", "Type", "Required", "Description"],
        [
            ["title", "Text", "Yes", "Post title"],
            ["slug", "Text", "Yes", "URL: /blog/{slug}"],
            ["topic", "Option or Text", "No", "Category/tag"],
            ["author", "Text", "No", "Author name"],
            ["published_at", "Date/time", "Yes", "Publication date"],
            ["featured_image", "Asset", "No", "Hero image"],
            ["excerpt", "Textarea", "No", "Listing summary"],
            ["body", "Richtext", "Yes", "Article content"],
            ["seo_*", "Various", "No", "SEO metadata"],
        ],
        [1.3, 1.0, 0.7, 3.5],
    )

    doc.add_heading("5.6 Reusable Section Bloks", level=2)
    add_table(doc,
        ["Blok", "Key Fields", "Notes"],
        [
            ["section_hero", "eyebrow, headline, subheadline, image, cta_label, cta_link", "Primary hero"],
            ["section_richtext", "content (richtext)", "Free-form prose"],
            ["section_two_column", "left_blocks, right_blocks", "Split layout"],
            ["section_faq", "items[{question, answer}]", "Accordion FAQ"],
            ["section_testimonials", "source (manual|trustpilot), items[]", "Trustpilot via API optional"],
            ["section_cta", "headline, body, button_label, button_link", "Conversion block"],
            ["section_image_gallery", "images[]", "Media grid"],
            ["section_video", "youtube_id or asset", "Lazy-loaded embed"],
            ["course_price_block", "zenler_course_id, layout_variant, show_compare", "Fetches live geo price"],
            ["course_finder", "filters_config", "Dynamic catalog from API"],
            ["schema_organization", "name, url, logo, sameAs[]", "Site-wide schema"],
            ["schema_course", "name, description, provider, offers", "Course JSON-LD"],
        ],
        [1.8, 2.5, 2.2],
    )

    doc.add_heading("5.7 Global: header (Multi-Level Navigation)", level=2)
    doc.add_paragraph(
        "The header blok must support unlimited menu depth to replace Zenler navigation limitations."
    )
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["logo", "Asset", "Site logo"],
            ["nav_items", "Blocks: nav_item (recursive)", "Top-level navigation"],
            ["nav_item.label", "Text", "Menu label"],
            ["nav_item.link", "Link", "Internal or external URL"],
            ["nav_item.children", "Blocks: nav_item", "Nested submenu items"],
            ["cta_button", "Block: link_button", "Optional enroll/contact CTA"],
        ],
        [1.5, 1.5, 3.5],
    )

    doc.add_heading("5.8 Storyblok Webhooks", level=2)
    add_table(doc,
        ["Event", "Target", "Action"],
        [
            ["story.published", "POST https://vls-online.com/api/revalidate", "Revalidate affected paths/tags"],
            ["story.unpublished", "POST /api/revalidate", "Revalidate or return 404 on next request"],
            ["story.deleted", "POST /api/revalidate", "Purge cache for slug"],
        ],
        [1.5, 2.5, 2.5],
    )
    add_code(doc, """
POST /api/revalidate
Headers: Authorization: Bearer {REVALIDATE_SECRET}
Body: {
  "full_slug": "courses/fa1",
  "action": "published"
}

Response: { "revalidated": true, "paths": ["/en/courses/fa1"] }
""")

    doc.add_heading("5.9 Editorial Roles", level=2)
    add_table(doc,
        ["Role", "Permissions"],
        [
            ["Editor", "Create/edit/publish pages, blog, landing; edit course content except schema/globals"],
            ["Admin", "All Editor permissions + globals (header/footer) + asset management"],
            ["Developer", "Blok schema changes only; no routine content editing"],
        ],
        [1.5, 5.0],
    )

    # --- 6 Next.js ---
    doc.add_heading("6. Next.js Frontend Specification", level=1)

    doc.add_heading("6.1 Technology Stack", level=2)
    add_table(doc,
        ["Package", "Version", "Purpose"],
        [
            ["next", "15.x", "App Router, RSC, ISR"],
            ["react", "19.x", "UI library"],
            ["@storyblok/react", "latest", "Storyblok bridge and components"],
            ["next-intl", "latest", "Locale routing and translations"],
            ["tailwindcss", "4.x", "Styling (match VLS brand)"],
            ["stripe", "latest", "Client-side redirect to Checkout"],
            ["zod", "latest", "Runtime validation for API responses"],
        ],
        [2.0, 1.0, 3.5],
    )

    doc.add_heading("6.2 Repository Structure", level=2)
    add_code(doc, """
vls-web/
  app/
    [locale]/
      layout.tsx
      page.tsx                    # home
      courses/[slug]/page.tsx
      blog/page.tsx
      blog/[slug]/page.tsx
      [...slug]/page.tsx            # flexible pages
    checkout/[courseSlug]/page.tsx
    payment-success/page.tsx
    api/revalidate/route.ts
    sitemap.ts
    robots.ts
  components/
    storyblok/                    # blok React components
    layout/Header.tsx
    layout/Footer.tsx
    pricing/CoursePriceBlock.tsx
  lib/
    storyblok.ts
    api.ts
    geo.ts
    schema.ts
  messages/                       # next-intl JSON
    en.json
    ar.json
  middleware.ts                   # locale detection + geo header
  next.config.ts
""")

    doc.add_heading("6.3 Rendering Strategy", level=2)
    add_table(doc,
        ["Page Type", "Strategy", "Revalidation", "Notes"],
        [
            ["Marketing pages", "ISR", "60s or webhook", "fetchStoryblok with cv=published"],
            ["Course pages", "ISR + dynamic islands", "60s + price API no-cache", "Price fetched server-side per request"],
            ["Blog posts", "ISR", "300s or webhook", "Static body, cached shell"],
            ["Checkout", "SSR dynamic", "none", "Always fresh Stripe session"],
            ["Draft preview", "SSR", "none", "Storyblok draft mode with bridge"],
        ],
        [1.5, 1.5, 1.5, 2.0],
    )

    doc.add_heading("6.4 Storyblok Integration", level=2)
    add_bullets(doc, [
        "Use Storyblok Content Delivery API v2 for published content.",
        "Enable Storyblok Bridge in draft/preview mode for Visual Editor live preview.",
        "Map each blok schema name to a React component in components/storyblok/.",
        "Use storyblokEditable() for in-context editing in preview.",
        "Configure preview URL: https://preview.vls-online.com/{path}?_storyblok={story_id}",
    ])

    doc.add_heading("6.5 Dynamic Course Price Component", level=2)
    add_code(doc, """
// Server Component: CoursePriceBlock.tsx
async function CoursePriceBlock({ blok, locale }) {
  const country = headers().get('x-vls-country') ?? 'GB';
  const price = await fetch(
    `${API_URL}/api/v1/prices/${blok.zenler_course_id}?country=${country}&locale=${locale}`,
    { next: { revalidate: 0 } }
  ).then(r => r.json());

  return (
    <PricingCard
      amount={price.amount}
      currency={price.currency}
      compareAt={price.compareAt}
      checkoutUrl={`/checkout/${price.courseSlug}?option=${price.paymentOptionId}`}
    />
  );
}
""")

    doc.add_heading("6.6 SEO and Structured Data", level=2)
    add_bullets(doc, [
        "Every page exports generateMetadata() from Storyblok SEO fields.",
        "Course pages emit Course and Offer JSON-LD with localized price from API.",
        "Organization schema on all pages via site-config global.",
        "BreadcrumbList schema on course and blog pages.",
        "Canonical URLs include locale prefix: https://vls-online.com/en/courses/fa1",
        "hreflang alternates for each supported locale.",
        "Automatic sitemap.xml from Storyblok stories + course API.",
    ])

    doc.add_heading("6.7 Performance Requirements", level=2)
    add_bullets(doc, [
        "Use next/image for all Storyblok assets with defined width/height.",
        "Lazy load below-fold section bloks and video embeds.",
        "Self-host Poppins/Inter fonts via next/font.",
        "Target JavaScript bundle < 150KB gzipped on course landing pages.",
        "Enable Vercel Speed Insights and Lighthouse CI in pipeline.",
        "Cloudflare cache rules for /_next/static/* (1 year) and images (30 days).",
    ])

    # --- 7 Cloudflare ---
    doc.add_heading("7. Cloudflare Configuration Specification", level=1)

    doc.add_heading("7.1 DNS Records", level=2)
    add_table(doc,
        ["Type", "Name", "Target", "Proxy"],
        [
            ["CNAME", "vls-online.com", "cname.vercel-dns.com", "Proxied (orange cloud)"],
            ["CNAME", "www", "vls-online.com", "Proxied"],
            ["CNAME", "learn", "Zenler-provided CNAME", "Proxied or DNS only per Zenler docs"],
            ["CNAME", "api", "cname.vercel-dns.com", "Proxied"],
        ],
        [0.8, 1.2, 2.5, 1.0],
    )

    doc.add_heading("7.2 Cache Rules", level=2)
    add_table(doc,
        ["Rule", "Match", "Cache TTL", "Notes"],
        [
            ["Static assets", "/_next/static/*", "31536000", "Immutable hashed assets"],
            ["Images", "/_next/image*", "2592000", "Optimized images"],
            ["ISR pages", "HTML documents on marketing paths", "Respect s-maxage from origin", "Do not override short ISR"],
            ["API", "api.vls-online.com/*", "Bypass", "Never cache authenticated or checkout APIs"],
        ],
        [1.2, 2.0, 1.3, 2.0],
    )

    doc.add_heading("7.3 Security", level=2)
    add_bullets(doc, [
        "Enable Cloudflare WAF managed ruleset on vls-online.com.",
        "Rate limit POST /api/v1/checkout and form endpoints: 10 req/min per IP.",
        "Turnstile on contact and lead forms (existing integration retained).",
        "Strict TLS 1.2+; HSTS max-age 31536000 includeSubDomains.",
        "Block direct origin access; allow only Cloudflare IP ranges to Vercel (if supported).",
    ])

    doc.add_heading("7.4 Geo Country Header", level=2)
    doc.add_paragraph(
        "Cloudflare automatically provides CF-IPCountry on proxied requests. A Worker or Transform Rule "
        "must copy this to x-vls-country for the Next.js origin."
    )
    add_code(doc, """
// Transform Rule (alternative to Worker)
// Modify Request Header: set x-vls-country = cf.ipcountry
""")

    # --- 8 Backend API ---
    doc.add_heading("8. Backend API Specification", level=1)
    doc.add_paragraph(
        "The backend evolves from cms-v2 Express API. Marketing content endpoints (publish-*) are "
        "deprecated. Commerce, sync, and integration endpoints are retained and extended."
    )

    doc.add_heading("8.1 Service Deployment", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Hostname", "api.vls-online.com"],
            ["Runtime", "Node.js 20 on Vercel serverless or dedicated Node service"],
            ["Codebase", "cms-v2/src/ refactored into vls-api/"],
            ["Database", "Neon PostgreSQL (existing)"],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("8.2 New API Endpoints", level=2)
    add_table(doc,
        ["Method", "Path", "Auth", "Description"],
        [
            ["GET", "/api/v1/prices/:courseSlug", "Public", "Resolve geo price by CF country header or ?country=XX"],
            ["GET", "/api/v1/courses", "Public", "Course catalog for course finder (from Zenler sync)"],
            ["GET", "/api/v1/courses/:slug", "Public", "Single course metadata"],
            ["POST", "/api/v1/checkout/session", "Public", "Create Stripe Checkout Session (geo-aware)"],
            ["GET", "/api/v1/checkout/status", "Public", "Payment status by session_id"],
            ["POST", "/api/v1/forms/contact", "Public + Turnstile", "Contact form submission"],
            ["POST", "/api/v1/webhooks/stripe", "Stripe signature", "Payment webhook (existing, extended)"],
            ["POST", "/api/v1/webhooks/storyblok", "Storyblok token", "Optional content sync events"],
            ["POST", "/api/v1/admin/zenler/sync", "Admin JWT", "Trigger course catalog sync"],
            ["POST", "/api/v1/admin/prices", "Admin JWT", "Manage regional price matrix"],
        ],
        [0.7, 2.5, 1.0, 2.3],
    )

    doc.add_heading("8.3 GET /api/v1/prices/:courseSlug", level=2)
    add_code(doc, """
Request:
  GET /api/v1/prices/fa1?country=AE&locale=en
  Header: x-vls-country: AE (optional override)

Response 200:
{
  "courseSlug": "fa1",
  "zenlerCourseId": "12345",
  "country": "AE",
  "region": "GCC",
  "currency": "AED",
  "amount": 1299.00,
  "compareAt": 1499.00,
  "paymentOptionId": 42,
  "stripePriceId": "price_xxx",
  "formatted": "AED 1,299"
}

Response 404: course or price not found for region (fall back to default region in client)
""")

    doc.add_heading("8.4 POST /api/v1/checkout/session", level=2)
    add_code(doc, """
Request:
{
  "courseSlug": "fa1",
  "paymentOptionId": 42,
  "country": "AE",
  "studentEmail": "student@example.com",
  "studentName": "Jane Doe",
  "locale": "en"
}

Response 200:
{ "checkoutUrl": "https://checkout.stripe.com/c/pay/..." }

Logic:
  1. Resolve regional price for country
  2. Create payment_orders row
  3. Create Stripe Checkout Session using stripePriceId (preferred) or price_data
  4. Attach zenlerCourseId, orderId, paymentOptionId in metadata
""")

    doc.add_heading("8.5 Stripe Webhook Extension", level=2)
    doc.add_paragraph(
        "Extend existing stripeWebhookHandler to call Zenler enrollment after markPaymentOrderPaid."
    )
    add_numbered(doc, [
        "Verify Stripe signature (existing: verifyStripeWebhook).",
        "Handle checkout.session.completed (existing).",
        "Mark order paid in payment_orders (existing).",
        "NEW: Call enrollStudentInZenler({ zenlerCourseId, email, name, orderId }).",
        "NEW: Record enrollment status in payment_orders or new zenler_enrollments table.",
        "Send student confirmation email with learn.vls-online.com login link (existing + updated template).",
        "Send admin notification (existing).",
        "On enrollment failure: retry 3x with exponential backoff; alert admin; do not refund automatically.",
    ])

    doc.add_heading("8.6 Zenler Enrollment Service (New)", level=2)
    add_code(doc, """
// services/zenlerEnrollment.ts
export async function enrollStudentInZenler(input: {
  zenlerCourseId: string;
  email: string;
  name?: string;
  orderId: number;
}): Promise<{ success: boolean; zenlerUserId?: string; error?: string }>

Implementation notes:
  - Use Zenler API (confirm exact endpoint with Zenler support/docs)
  - Idempotent: check if user already enrolled before creating duplicate
  - Store enrollment result linked to payment_order_id
  - Required env: ZENLER_API_KEY, ZENLER_ACCOUNT_NAME
""")

    doc.add_heading("8.7 Deprecated Endpoints (Post-Migration)", level=2)
    add_bullets(doc, [
        "GET /api/publish-banner",
        "GET /api/publish-header",
        "GET /api/publish-course-prices",
        "GET /api/publish-bpp-books",
        "GET /api/publish-testimonials-components",
        "All cms-v2 /api/content snippet publish routes used for HTML injects.",
        "cms-v2 admin SPA (dist-client).",
    ])

    # --- 9 Database ---
    doc.add_heading("9. Database Schema Changes", level=1)

    doc.add_heading("9.1 Existing Tables (Retained)", level=2)
    add_bullets(doc, [
        "courses - synced from Zenler API",
        "course_payment_cards - payment options per course",
        "payment_orders - checkout and payment lifecycle",
        "books, book_discount_codes - BPP book fulfilment",
        "cms_activity_logs - audit trail",
    ])

    doc.add_heading("9.2 New Table: price_regions", level=2)
    add_code(doc, """
CREATE TABLE price_regions (
  id              SERIAL PRIMARY KEY,
  code            VARCHAR(16) NOT NULL UNIQUE,  -- e.g. UK, US, EU, GCC, DEFAULT
  name            TEXT NOT NULL,
  default_currency VARCHAR(8) NOT NULL,
  countries       TEXT[] NOT NULL,               -- ISO 3166-1 alpha-2 codes
  is_active       BOOLEAN NOT NULL DEFAULT TRUE,
  sort_order      INT NOT NULL DEFAULT 0,
  created_at      TIMESTAMPTZ NOT NULL DEFAULT NOW()
);
""")

    doc.add_heading("9.3 New Table: course_regional_prices", level=2)
    add_code(doc, """
CREATE TABLE course_regional_prices (
  id                  SERIAL PRIMARY KEY,
  course_id           INT NOT NULL REFERENCES courses(id) ON DELETE CASCADE,
  payment_option_id   INT REFERENCES course_payment_cards(id) ON DELETE SET NULL,
  region_code         VARCHAR(16) NOT NULL REFERENCES price_regions(code),
  currency            VARCHAR(8) NOT NULL,
  amount              NUMERIC(10,2) NOT NULL,
  compare_at_amount   NUMERIC(10,2),
  stripe_price_id     TEXT,                      -- preferred for Checkout
  is_active           BOOLEAN NOT NULL DEFAULT TRUE,
  created_at          TIMESTAMPTZ NOT NULL DEFAULT NOW(),
  updated_at          TIMESTAMPTZ NOT NULL DEFAULT NOW(),
  UNIQUE (course_id, region_code, payment_option_id)
);

CREATE INDEX idx_course_regional_prices_lookup
  ON course_regional_prices (course_id, region_code) WHERE is_active = TRUE;
""")

    doc.add_heading("9.4 New Table: zenler_enrollments", level=2)
    add_code(doc, """
CREATE TABLE zenler_enrollments (
  id                  SERIAL PRIMARY KEY,
  payment_order_id    INT NOT NULL REFERENCES payment_orders(id),
  zenler_course_id    TEXT NOT NULL,
  student_email       TEXT NOT NULL,
  zenler_user_id      TEXT,
  status              TEXT NOT NULL DEFAULT 'pending',  -- pending|success|failed
  attempts            INT NOT NULL DEFAULT 0,
  last_error          TEXT,
  enrolled_at         TIMESTAMPTZ,
  created_at          TIMESTAMPTZ NOT NULL DEFAULT NOW()
);
""")

    doc.add_heading("9.5 Seed Data: Initial Regions", level=2)
    add_table(doc,
        ["Region Code", "Currency", "Example Countries"],
        [
            ["UK", "GBP", "GB"],
            ["EU", "EUR", "DE, FR, NL, IE, ES, IT, ..."],
            ["US", "USD", "US, CA (configurable)"],
            ["GCC", "AED", "AE, SA, QA, KW, BH, OM"],
            ["DEFAULT", "USD", "Fallback for unmapped countries"],
        ],
        [1.5, 1.2, 3.8],
    )

    # --- 10 Stripe ---
    doc.add_heading("10. Stripe Geo-Pricing Specification", level=1)

    doc.add_heading("10.1 Stripe Object Model", level=2)
    add_table(doc,
        ["Stripe Object", "Mapping", "Notes"],
        [
            ["Product", "One per course (or payment card)", "Named e.g. VLS FA1 Full Course"],
            ["Price", "One per region/currency/plan", "Recurring or one-time as required"],
            ["Checkout Session", "Created per purchase", "Uses stripe_price_id from course_regional_prices"],
            ["Webhook", "checkout.session.completed", "Triggers enrollment"],
        ],
        [1.5, 2.5, 2.5],
    )

    doc.add_heading("10.2 Country to Region Resolution", level=2)
    add_code(doc, """
function resolveRegion(countryCode: string): string {
  // 1. Look up country in price_regions.countries arrays
  // 2. If multiple matches, use lowest sort_order
  // 3. If no match, return 'DEFAULT'
}
""")

    doc.add_heading("10.3 Admin Price Management", level=2)
    doc.add_paragraph(
        "Phase 1: Admin UI in slimmed cms-v2 admin or standalone internal tool. "
        "Phase 2: Storyblok plugin or external spreadsheet sync. "
        "Editors must NOT enter Stripe Price IDs in Storyblok; they select course only."
    )

    # --- 11 Zenler ---
    doc.add_heading("11. Zenler LMS Integration Specification", level=1)

    doc.add_heading("11.1 Zenler Scope (Post-Migration)", level=2)
    add_table(doc,
        ["Function", "Platform", "Notes"],
        [
            ["Marketing pages", "Next.js / Storyblok", "Removed from Zenler public site"],
            ["Blog", "Next.js / Storyblok", "Removed from Zenler"],
            ["Video hosting", "Zenler", "No migration"],
            ["Curriculum structure", "Zenler", "Modules, lessons, drip"],
            ["Student login", "Zenler on learn.vls-online.com", "Custom domain"],
            ["Progress tracking", "Zenler", "Unchanged"],
            ["Certificates", "Zenler", "If enabled"],
            ["Course catalog sync", "Backend API to Neon", "Existing zenlerCourseService"],
            ["Enrollment after payment", "Backend API", "New zenlerEnrollment service"],
        ],
        [2.0, 2.0, 2.5],
    )

    doc.add_heading("11.2 Zenler API Usage", level=2)
    add_table(doc,
        ["Operation", "Existing/New", "Endpoint Pattern"],
        [
            ["List courses", "Existing", "GET /api/v1/courses (paginated)"],
            ["Enroll user", "New - confirm with Zenler", "TBD - Zenler enrollment API"],
            ["Create user", "New - if required", "TBD - Zenler user API"],
        ],
        [1.5, 1.5, 3.5],
    )
    doc.add_paragraph(
        "ACTION ITEM: Confirm Zenler enrollment API capabilities and rate limits with Zenler support "
        "before Phase 0 completion. If API enrollment is unavailable, fallback is manual admin enrollment "
        "with automated alert (unacceptable for production; must resolve in spike)."
    )

    doc.add_heading("11.3 Catalog Sync Schedule", level=2)
    add_bullets(doc, [
        "Scheduled sync: daily at 02:00 UTC via Vercel Cron or external scheduler.",
        "Manual sync: POST /api/v1/admin/zenler/sync (admin only).",
        "Sync updates courses table: zenlerCourseId, name, slug, category, level, zenlerUrl.",
        "zenlerUrl field updated to learn.vls-online.com paths where applicable.",
    ])

    # --- 12 i18n ---
    doc.add_heading("12. Internationalization (i18n) Specification", level=1)

    doc.add_heading("12.1 Supported Locales (Phase 1)", level=2)
    add_table(doc,
        ["Locale", "Language", "URL Prefix", "Priority"],
        [
            ["en", "English", "/en/", "Launch"],
            ["ar", "Arabic", "/ar/", "Phase 4 (RTL layout required)"],
        ],
        [1.0, 1.5, 1.5, 2.5],
    )

    doc.add_heading("12.2 Storyblok Translation Strategy", level=2)
    doc.add_paragraph(
        "Use Storyblok field-level translation for shared page structure. Enable the Translatable "
        "Fields app. Each translatable field stores values per locale (en, ar)."
    )
    add_bullets(doc, [
        "Default locale: en",
        "Fallback: if ar field empty, fall back to en",
        "RTL: apply dir=rtl and mirrored layout for ar locale in Next.js layout",
        "Date/number formatting via next-intl",
        "Translated slugs: optional Phase 2 (/ar/courses/...); Phase 1 may share slug across locales",
    ])

    doc.add_heading("12.3 Locale Detection Middleware", level=2)
    add_code(doc, """
// middleware.ts
// 1. If path starts with /en or /ar, use that locale
// 2. Else redirect to preferred locale from Accept-Language or x-vls-country mapping
// 3. Do NOT auto-redirect logged-in LMS traffic (learn subdomain)
""")

    # --- 13 Security ---
    doc.add_heading("13. Security Specification", level=1)
    add_table(doc,
        ["Area", "Requirement"],
        [
            ["Secrets", "All keys in Vercel/Cloudflare env; never in Storyblok or client bundles"],
            ["Stripe webhooks", "Verify signature on raw body; idempotent order processing"],
            ["Admin API", "JWT auth with role checks (retain cms-v2 auth model)"],
            ["Revalidate endpoint", "Bearer token REVALIDATE_SECRET"],
            ["CORS", "Public read APIs: allow vls-online.com origins only"],
            ["PII", "Student email stored in payment_orders; GDPR retention policy applies"],
            ["CSP", "Strict Content-Security-Policy on Next.js; allow Stripe, Storyblok, Turnstile domains"],
        ],
        [1.5, 5.0],
    )

    # --- 14 Environment Variables ---
    doc.add_heading("14. Environment Variables", level=1)
    add_table(doc,
        ["Variable", "Service", "Description"],
        [
            ["STORYBLOK_TOKEN", "Next.js (server)", "Content Delivery API token (published)"],
            ["STORYBLOK_PREVIEW_TOKEN", "Next.js (server)", "Preview/draft token"],
            ["REVALIDATE_SECRET", "Next.js", "Webhook auth for ISR revalidation"],
            ["NEXT_PUBLIC_SITE_URL", "Next.js", "https://vls-online.com"],
            ["API_URL", "Next.js", "https://api.vls-online.com"],
            ["STRIPE_SECRET_KEY", "API", "Stripe secret key"],
            ["STRIPE_WEBHOOK_SECRET", "API", "Webhook signing secret"],
            ["DATABASE_URL", "API", "Neon connection string"],
            ["ZENLER_API_KEY", "API", "Zenler API key"],
            ["ZENLER_ACCOUNT_NAME", "API", "Zenler account name"],
            ["TURNSTILE_SECRET_KEY", "API", "Cloudflare Turnstile"],
            ["NEXT_PUBLIC_TURNSTILE_SITE_KEY", "Next.js", "Turnstile site key"],
            ["JWT_SECRET", "API", "Admin authentication"],
        ],
        [2.2, 1.3, 3.0],
    )

    # --- 15 Migration ---
    doc.add_heading("15. Migration Plan", level=1)

    doc.add_heading("15.1 Phase Overview", level=2)
    add_table(doc,
        ["Phase", "Duration", "Deliverables"],
        [
            ["0 - Foundation", "2-3 weeks", "Storyblok space, Next.js scaffold, Cloudflare DNS, geo price schema, Zenler enrollment spike"],
            ["1 - Pilot course", "3-4 weeks", "Header/footer, one course page E2E with Stripe + enrollment, SEO baseline"],
            ["2 - Core marketing", "4-6 weeks", "Home, about, legal, course finder, redirect map"],
            ["3 - Blog", "2-3 weeks", "Blog migration script, listing + post templates, 301 redirects"],
            ["4 - i18n + regions", "3-4 weeks", "Arabic/RTL, full regional price matrix"],
            ["5 - Decommission cms-v2", "1-2 weeks", "Retire inject scripts, admin editors, publish endpoints"],
        ],
        [1.8, 1.2, 4.5],
    )

    doc.add_heading("15.2 Blog Migration Script Requirements", level=2)
    add_bullets(doc, [
        "Read posts from Neon vls-blog-posts table.",
        "Upload images to Storyblok asset manager (from cms_blog_assets BYTEA or exported files).",
        "Create blog_post stories with slug, title, body, published_at, topic.",
        "Rewrite internal links using compactBlogPathname logic from shared/blogUrls.ts.",
        "Generate redirect CSV for Cloudflare Bulk Redirects.",
    ])

    doc.add_heading("15.3 Parallel Run Strategy", level=2)
    add_numbered(doc, [
        "Deploy Next.js to production with feature-flagged routes.",
        "Migrate one course (pilot) and validate payments end-to-end.",
        "Enable Cloudflare redirects for migrated paths only.",
        "Monitor Search Console, Stripe dashboard, enrollment success rate for 2 weeks.",
        "Migrate remaining pages in batches by traffic priority.",
        "Cut Zenler marketing domain mapping; point learn.vls-online.com only.",
    ])

    # --- 16 Testing ---
    doc.add_heading("16. Testing and QA Specification", level=1)
    add_table(doc,
        ["Test Type", "Scope", "Tool"],
        [
            ["Unit", "Price resolution, region mapping, slug helpers", "Vitest"],
            ["Integration", "Stripe webhook, Zenler enrollment mock", "Vitest + MSW"],
            ["E2E", "Course page load, checkout flow, payment success", "Playwright"],
            ["Visual", "Storyblok blok rendering regression", "Chromatic or Percy"],
            ["Performance", "Lighthouse CI on course pages", "lighthouse-ci"],
            ["SEO", "Schema validation, sitemap, hreflang", "Manual + Google Rich Results Test"],
            ["Content QA", "Editor publish workflow in Storyblok preview", "Manual checklist"],
        ],
        [1.2, 3.0, 2.3],
    )

    # --- 17 Monitoring ---
    doc.add_heading("17. Monitoring and Observability", level=1)
    add_bullets(doc, [
        "Vercel Analytics and Speed Insights on Next.js.",
        "Stripe Dashboard for payment failures and dispute monitoring.",
        "API structured logging for webhook and enrollment events (existing console.error patterns extended).",
        "Alert on zenler_enrollments.status=failed after 3 retries (email to office@vls-online.com).",
        "Uptime monitoring on vls-online.com, learn.vls-online.com, api.vls-online.com.",
        "Cloudflare analytics for cache hit ratio and WAF events.",
    ])

    # --- 18 Risks ---
    doc.add_heading("18. Risks and Mitigations", level=1)
    add_table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["Zenler enrollment API unavailable", "Medium", "High", "Phase 0 spike; escalate with Zenler support"],
            ["SEO ranking drop during migration", "Medium", "High", "301 redirects; keep slugs; monitor GSC"],
            ["Editor publishes broken layout", "Medium", "Medium", "Constrained blok set; preview required; training"],
            ["Geo price misconfiguration", "Low", "High", "Admin review workflow; Stripe test mode validation"],
            ["Storyblok cost overrun", "Low", "Medium", "Right-size plan; limit seats"],
            ["Dual content during migration", "High", "Medium", "Clear ownership per URL; redirect registry"],
        ],
        [2.5, 1.0, 1.0, 2.0],
    )

    # --- 19 Appendices ---
    doc.add_heading("Appendix A: Storyblok Blok Component Registry", level=1)
    add_table(doc,
        ["Blok Name", "React Component", "Render Type"],
        [
            ["page", "Page.tsx", "Static/ISR"],
            ["course_page", "CoursePage.tsx", "ISR + dynamic price"],
            ["blog_post", "BlogPost.tsx", "ISR"],
            ["section_hero", "SectionHero.tsx", "Static"],
            ["section_faq", "SectionFaq.tsx", "Static"],
            ["section_testimonials", "SectionTestimonials.tsx", "Static/API"],
            ["section_cta", "SectionCta.tsx", "Static"],
            ["course_price_block", "CoursePriceBlock.tsx", "Dynamic server"],
            ["course_finder", "CourseFinder.tsx", "Dynamic server"],
            ["header (global)", "Header.tsx", "ISR"],
            ["footer (global)", "Footer.tsx", "ISR"],
        ],
        [2.0, 2.5, 1.5],
    )

    doc.add_heading("Appendix B: Payment Order State Machine", level=1)
    add_code(doc, """
pending --> checkout_created (Stripe session created)
checkout_created --> paid (webhook checkout.session.completed)
paid --> enrolling (Zenler enrollment started)
enrolling --> enrollment_success (student can access learn.vls-online.com)
enrolling --> enrollment_failed (retry + admin alert; payment NOT auto-refunded)
paid --> emails_sent (confirmation emails dispatched)
""")

    doc.add_heading("Appendix C: Content Editor Quick Reference", level=1)
    add_table(doc,
        ["Task", "Steps"],
        [
            ["Create campaign landing page", "Storyblok > landing > New story > Add section bloks > Publish"],
            ["Update course FAQ", "Storyblok > courses > {course} > Edit faq blok > Publish"],
            ["Change header menu", "Storyblok > global > header > Edit nav_items > Publish"],
            ["Publish blog post", "Storyblok > blog > posts > New > Write > Set published_at > Publish"],
            ["Change course price", "Admin price tool or Stripe (NOT Storyblok)"],
        ],
        [2.5, 4.0],
    )

    doc.add_heading("Appendix D: Glossary", level=1)
    add_table(doc,
        ["Term", "Definition"],
        [
            ["Blok", "Storyblok content component (schema block)"],
            ["ISR", "Incremental Static Regeneration (Next.js)"],
            ["LMS", "Learning Management System (Zenler)"],
            ["RSC", "React Server Components"],
            ["Geo pricing", "Displaying and charging region-appropriate prices"],
            ["Inject script", "Legacy cms-v2 HTML/JS pasted into Zenler pages (to be retired)"],
        ],
        [1.5, 5.0],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build()
